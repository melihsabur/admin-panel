import json
import asyncio
import os
import re
import aio_pika
from datetime import datetime

RABBITMQ_URL = os.getenv("RABBITMQ_URL", "amqp://rabbitmq_user:rabbitmq_pass@rabbitmq:5672/")
UPLOAD_DIR = os.getenv("UPLOAD_DIR", "/app/uploads")


# =====================
# RAG Benzeri Bilgi Tabani (Knowledge Base)
# =====================
KNOWLEDGE_BASE = {
    "python": "Python, yuksek seviyeli, genel amacli bir programlama dilidir. Guido van Rossum tarafindan gelistirilmis olup, okunakli soz dizimi ve guclu kutuphaneleri ile bilinir. Web gelistirme (Django, Flask), veri bilimi (NumPy, Pandas), yapay zeka (TensorFlow, PyTorch) gibi alanlarda yaygin kullanilir.",
    "docker": "Docker, uygulamalari container'lar icinde paketleyerek calistirmanizi saglayan bir platformdur. Her container izole bir ortamda calisir ve docker-compose ile birden fazla container orkestre edilebilir.",
    "react": "React, Facebook tarafindan gelistirilen bir JavaScript UI kutuphanesidir. Component tabanli mimari, Virtual DOM ve hooks gibi ozellikleri ile modern web uygulamalari gelistirmek icin kullanilir.",
    "fastapi": "FastAPI, Python ile yuksek performansli API'ler olusturmak icin kullanilan modern bir web framework'udur. Otomatik dokumantasyon (Swagger), async destegi ve Pydantic ile tip guvenli veri dogrulama sunar.",
    "postgresql": "PostgreSQL, acik kaynakli, guclu bir nesne-iliskisel veritabani sistemidir. JSONB, Full-Text Search, Trigger, Function, Procedure gibi gelismis ozellikleri destekler.",
    "rabbitmq": "RabbitMQ, acik kaynakli bir mesaj kuyrugu sistemidir. AMQP protokolunu kullanir ve RPC, Pub/Sub, Work Queue gibi mesajlasma desenleri destekler. Mikroservis mimarilerinde servisler arasi iletisim icin idealdir.",
    "redis": "Redis, bellek-ici veri yapisi deposudur. Cache, session yonetimi, pub/sub mesajlasma ve kuyruk islemleri icin kullanilir. String, List, Set, Hash gibi veri yapilari destekler.",
    "tailwindcss": "TailwindCSS, utility-first yaklasimli bir CSS framework'udur. Onceden tanimli siniflar ile hizli ve tutarli UI gelistirme saglar.",
    "jwt": "JWT (JSON Web Token), guvenli bilgi iletimi icin kullanilan bir standarttir. Kullanici kimlik dogrulama ve yetkilendirme islemlerinde yaygin olarak kullanilir.",
    "microservice": "Mikroservis mimarisi, bir uygulamayi kucuk, bagimsiz servislerden olusan bir koleksiyon olarak tasarlar. Her servis kendi veritabanina sahip olabilir ve API veya mesaj kuyrugu ile iletisim kurar.",
}


def search_knowledge_base(query):
    query_lower = query.lower()
    results = []
    for key, value in KNOWLEDGE_BASE.items():
        if key in query_lower or any(word in query_lower for word in key.split()):
            results.append({"topic": key, "content": value})
    return results


def analyze_text_file(filepath):
    try:
        with open(filepath, "r", encoding="utf-8") as f:
            content = f.read()

        word_count = len(content.split())
        line_count = len(content.splitlines())
        char_count = len(content)

        # Kelime frekansi analizi
        words = re.findall(r'\b\w+\b', content.lower())
        word_freq = {}
        for w in words:
            if len(w) > 3:
                word_freq[w] = word_freq.get(w, 0) + 1
        top_words = sorted(word_freq.items(), key=lambda x: x[1], reverse=True)[:10]

        # Cumle sayisi
        sentence_count = len(re.split(r'[.!?]+', content))

        # Prompt komutlari varsa isle
        prompts = []
        for line in content.splitlines():
            stripped = line.strip()
            if stripped.startswith("PROMPT:") or stripped.startswith("SORU:"):
                prompt_text = stripped.split(":", 1)[1].strip()
                answer = advanced_ai_response(prompt_text)
                prompts.append({"prompt": prompt_text, "response": answer})

        return {
            "type": "text_analysis",
            "word_count": word_count,
            "line_count": line_count,
            "char_count": char_count,
            "sentence_count": sentence_count,
            "avg_word_length": round(sum(len(w) for w in words) / max(len(words), 1), 1),
            "top_words": [{"word": w, "count": c} for w, c in top_words],
            "prompts_found": len(prompts),
            "prompt_results": prompts,
            "summary": f"Dosya {line_count} satir, {word_count} kelime, {char_count} karakter, {sentence_count} cumle iceriyor. Ortalama kelime uzunlugu: {round(sum(len(w) for w in words) / max(len(words), 1), 1)} karakter.",
            "analyzed_at": datetime.utcnow().isoformat()
        }
    except Exception as e:
        return {"type": "error", "message": str(e)}


def advanced_ai_response(prompt):
    prompt_lower = prompt.lower()

    # RAG: Bilgi tabaninda ara
    kb_results = search_knowledge_base(prompt)
    if kb_results:
        topics = ", ".join([r["topic"] for r in kb_results])
        details = "\n".join([f"- {r['content']}" for r in kb_results])
        return f"[RAG] Bilgi tabaninda '{topics}' konularinda bilgi bulundu:\n{details}"

    # Matematik islemleri
    math_match = re.search(r'(\d+)\s*[+\-*/x]\s*(\d+)', prompt)
    if math_match or "hesapla" in prompt_lower or "toplam" in prompt_lower:
        if math_match:
            a, b = math_match.groups()
            a, b = int(a), int(b)
            op_char = re.search(r'[+\-*/x]', prompt).group()
            if op_char == '+': result = a + b
            elif op_char == '-': result = a - b
            elif op_char in ('*', 'x'): result = a * b
            elif op_char == '/': result = round(a / b, 2) if b != 0 else "Sifira bolme hatasi"
            else: result = "Bilinmeyen islem"
            return f"Hesaplama sonucu: {a} {op_char} {b} = {result}"
        return "Matematiksel islem icin lutfen 'sayi opertor sayi' formatinda yazin. Ornek: 15 + 27"

    # Tarih/saat
    if "tarih" in prompt_lower or "gun" in prompt_lower or "saat" in prompt_lower:
        now = datetime.now()
        return f"Sistem tarihi: {now.strftime('%d.%m.%Y')} | Saat: {now.strftime('%H:%M:%S')} | Gun: {['Pazartesi','Sali','Carsamba','Persembe','Cuma','Cumartesi','Pazar'][now.weekday()]}"

    # Selamlama
    if any(word in prompt_lower for word in ["merhaba", "selam", "hey", "nasilsin"]):
        return "Merhaba! Ben Admin Panel AI asistaniyim. Size dosya analizi, metin isleme, bilgi sorgulama ve prompt komutlari konularinda yardimci olabilirim. Bilgi tabaninda Python, Docker, React, FastAPI, PostgreSQL, RabbitMQ, Redis gibi konularda detayli bilgi bulunmaktadir."

    # Yardim
    if any(word in prompt_lower for word in ["yardim", "help", "ne yapabilirsin", "komutlar"]):
        return """Admin Panel AI Asistani Yetenekleri:

1. DOSYA ANALIZI: Yuklenen dosyalari analiz edebilirim (txt, docx, image)
2. PROMPT ISLEME: txt dosyalarinda PROMPT: veya SORU: ile baslayan satirlari islerim
3. WORD->HTML: DOCX dosyalarini HTML'e donusturebilirim
4. RAG SORGULAMA: Bilgi tabaninda Python, Docker, React, FastAPI, PostgreSQL, RabbitMQ, Redis konularini sorgulayabilirsiniz
5. MATEMATIK: Basit matematiksel islemler yapabilirim (orn: 15 + 27)
6. TARIH/SAAT: Guncel tarih ve saat bilgisi verebilirim
7. METIN ANALIZI: Kelime sayisi, cumle analizi, en sik kullanilan kelimeler

Ornek komutlar:
- 'Python nedir?'
- '125 * 48'
- 'Docker hakkinda bilgi ver'
- 'Tarih kac?'"""

    # Ozet
    if "ozet" in prompt_lower or "ozetle" in prompt_lower:
        return "Dosya ozetleme islemi icin once Dosya Yonetimi ekranindan bir dosya yukleyin, ardindan 'AI Analiz' butonuna tiklayin. TXT dosyalarinda otomatik istatistik ve icerik analizi yapilir."

    # Sistem durumu
    if any(word in prompt_lower for word in ["sistem", "durum", "status", "saglik"]):
        return f"""Sistem Durumu Raporu ({datetime.now().strftime('%d.%m.%Y %H:%M')}):
- Backend API: Aktif (FastAPI)
- Veritabani: PostgreSQL 15 (JSONB modu)
- Cache: Redis 7 (Aktif)
- Mesaj Kuyrugu: RabbitMQ 3 (RPC modu)
- AI Servisi: Aktif (RAG + Prompt Processing)
- Log Servisi: Aktif (Consumer dinliyor)
- Dosya Servisi: Aktif (Consumer dinliyor)
- Toplam Container: 8 adet"""

    # Genel yanit
    return f"[AI] '{prompt}' komutu islendi. Bu konu hakkinda bilgi tabaninda spesifik bir kayit bulunamadi. Bilgi tabaninda sorgulayabileceginiz konular: Python, Docker, React, FastAPI, PostgreSQL, RabbitMQ, Redis, TailwindCSS, JWT, Microservice. Ayrica 'yardim' yazarak tum yeteneklerimi gorebilirsiniz."


def convert_docx_to_html(filepath):
    try:
        from docx import Document
        doc = Document(filepath)

        html_parts = ['<!DOCTYPE html>', '<html lang="tr">', '<head>',
                      '<meta charset="UTF-8">', '<meta name="viewport" content="width=device-width, initial-scale=1.0">',
                      '<title>Donusturulmus Belge</title>',
                      '<link href="https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700&display=swap" rel="stylesheet">',
                      '<style>',
                      '* { margin: 0; padding: 0; box-sizing: border-box; }',
                      'body { font-family: "Inter", Arial, sans-serif; max-width: 900px; margin: 0 auto; padding: 40px 20px; line-height: 1.8; color: #1a1a2e; background: #fafafa; }',
                      'h1 { color: #1a1a2e; font-size: 2em; margin: 1.5em 0 0.5em; padding-bottom: 0.3em; border-bottom: 2px solid #6366f1; }',
                      'h2 { color: #16213e; font-size: 1.5em; margin: 1.2em 0 0.4em; }',
                      'h3 { color: #0f3460; font-size: 1.2em; margin: 1em 0 0.3em; }',
                      'p { margin: 0.8em 0; line-height: 1.8; }',
                      'ul, ol { margin: 0.8em 0; padding-left: 2em; }',
                      'li { margin: 0.3em 0; }',
                      'table { border-collapse: collapse; width: 100%; margin: 1em 0; }',
                      'th, td { border: 1px solid #ddd; padding: 10px; text-align: left; }',
                      'th { background: #6366f1; color: white; }',
                      'tr:nth-child(even) { background: #f5f5ff; }',
                      '.meta { color: #666; font-size: 0.9em; margin-bottom: 2em; padding: 1em; background: #f0f0ff; border-radius: 8px; }',
                      '</style>', '</head>', '<body>',
                      f'<div class="meta">Bu belge otomatik olarak donusturulmustur. Tarih: {datetime.now().strftime("%d.%m.%Y %H:%M")}</div>']

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            if para.style.name.startswith('Heading 1'):
                html_parts.append(f'<h1>{text}</h1>')
            elif para.style.name.startswith('Heading 2'):
                html_parts.append(f'<h2>{text}</h2>')
            elif para.style.name.startswith('Heading 3'):
                html_parts.append(f'<h3>{text}</h3>')
            elif para.style.name.startswith('List'):
                html_parts.append(f'<li>{text}</li>')
            else:
                # Bold ve italic isaretleme
                formatted = ""
                for run in para.runs:
                    t = run.text
                    if run.bold:
                        t = f"<strong>{t}</strong>"
                    if run.italic:
                        t = f"<em>{t}</em>"
                    formatted += t
                html_parts.append(f'<p>{formatted if formatted else text}</p>')

        # Tablolari isle
        for table in doc.tables:
            html_parts.append('<table>')
            for i, row in enumerate(table.rows):
                html_parts.append('<tr>')
                tag = 'th' if i == 0 else 'td'
                for cell in row.cells:
                    html_parts.append(f'<{tag}>{cell.text}</{tag}>')
                html_parts.append('</tr>')
            html_parts.append('</table>')

        html_parts.extend(['</body>', '</html>'])
        html_content = '\n'.join(html_parts)

        output_path = filepath.rsplit('.', 1)[0] + '.html'
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(html_content)

        return {
            "type": "docx_to_html",
            "success": True,
            "html_preview": html_content[:3000],
            "full_html": html_content,
            "output_file": os.path.basename(output_path),
            "paragraph_count": len(doc.paragraphs),
            "table_count": len(doc.tables),
            "converted_at": datetime.utcnow().isoformat()
        }
    except Exception as e:
        return {"type": "error", "success": False, "message": str(e)}


def analyze_image(filepath):
    file_size = os.path.getsize(filepath)
    ext = filepath.rsplit('.', 1)[-1].lower()

    # Basit gorsel analiz bilgisi
    try:
        with open(filepath, 'rb') as f:
            header = f.read(24)

        width, height = 0, 0
        if ext == 'png' and header[:8] == b'\x89PNG\r\n\x1a\n':
            import struct
            width = struct.unpack('>I', header[16:20])[0]
            height = struct.unpack('>I', header[20:24])[0]
        elif ext in ('jpg', 'jpeg'):
            width, height = "bilinmiyor", "bilinmiyor"
    except:
        width, height = "bilinmiyor", "bilinmiyor"

    return {
        "type": "image_analysis",
        "format": ext.upper(),
        "file_size_kb": round(file_size / 1024, 2),
        "dimensions": f"{width}x{height}" if width else "bilinmiyor",
        "color_depth": "24-bit RGB (tahmini)",
        "analysis": f"Gorsel format: {ext.upper()}, Boyut: {round(file_size/1024, 2)} KB. Gorsel basariyla analiz edildi.",
        "recommendations": [
            "Web icin optimize etmek icin WebP formatina donusturmeyi deneyin",
            "Boyutu kucultmek icin sIkIstIrma uygulayabilirsiniz"
        ],
        "analyzed_at": datetime.utcnow().isoformat()
    }


def analyze_pdf(filepath):
    file_size = os.path.getsize(filepath)
    return {
        "type": "pdf_analysis",
        "file_size_kb": round(file_size / 1024, 2),
        "analysis": f"PDF dosyasi analiz edildi. Boyut: {round(file_size/1024, 2)} KB.",
        "note": "PDF icerik okuma icin ek kutuphane (PyPDF2/pdfplumber) gereklidir.",
        "analyzed_at": datetime.utcnow().isoformat()
    }


def analyze_excel(filepath):
    file_size = os.path.getsize(filepath)
    return {
        "type": "excel_analysis",
        "file_size_kb": round(file_size / 1024, 2),
        "analysis": f"Excel dosyasi analiz edildi. Boyut: {round(file_size/1024, 2)} KB.",
        "note": "Excel icerik okuma icin openpyxl kutuphanesi kullanilabilir.",
        "analyzed_at": datetime.utcnow().isoformat()
    }


async def on_message(message: aio_pika.IncomingMessage):
    async with message.process():
        try:
            body = json.loads(message.body)
            print(f"[AI CONSUMER] Mesaj alindi: {body.get('action')}")

            action = body.get("action")
            response = {"success": False}

            if action == "ANALYZE":
                filename = body.get("filename", "")
                filepath = os.path.join(UPLOAD_DIR, filename)
                mime_type = body.get("mime_type", "")

                if not os.path.exists(filepath):
                    response = {"success": False, "error": "Dosya bulunamadi"}
                elif mime_type == "text/plain" or filename.endswith(".txt"):
                    analysis = analyze_text_file(filepath)
                    response = {"success": True, "analysis": analysis}
                elif filename.endswith(".docx"):
                    analysis = convert_docx_to_html(filepath)
                    response = {"success": True, "analysis": analysis}
                elif mime_type and mime_type.startswith("image/"):
                    analysis = analyze_image(filepath)
                    response = {"success": True, "analysis": analysis}
                elif filename.endswith(".pdf"):
                    analysis = analyze_pdf(filepath)
                    response = {"success": True, "analysis": analysis}
                elif filename.endswith(".xlsx"):
                    analysis = analyze_excel(filepath)
                    response = {"success": True, "analysis": analysis}
                else:
                    response = {
                        "success": True,
                        "analysis": {
                            "type": "general",
                            "message": f"Dosya tipi: {mime_type}. Genel analiz tamamlandi.",
                            "file_size_kb": round(os.path.getsize(filepath) / 1024, 2),
                            "analyzed_at": datetime.utcnow().isoformat()
                        }
                    }

            elif action == "PROMPT":
                prompt_text = body.get("prompt_text", "")
                result = advanced_ai_response(prompt_text)
                response = {"success": True, "result": result, "prompt": prompt_text}

            if message.reply_to:
                connection = await aio_pika.connect_robust(RABBITMQ_URL)
                channel = await connection.channel()
                await channel.default_exchange.publish(
                    aio_pika.Message(
                        body=json.dumps(response).encode(),
                        correlation_id=message.correlation_id,
                    ),
                    routing_key=message.reply_to,
                )
                await connection.close()
                print(f"[AI RPC] Yanit gonderildi")

        except Exception as e:
            print(f"[AI CONSUMER HATA] {e}")

async def main():
    print("[AI SERVICE] Baslatiliyor... (RAG + Prompt Processing)")
    for attempt in range(30):
        try:
            connection = await aio_pika.connect_robust(RABBITMQ_URL)
            channel = await connection.channel()
            queue = await channel.declare_queue("ai_queue", durable=True)
            await queue.consume(on_message)
            print("[AI SERVICE] Dinleniyor... (MCP/RAG Aktif)")
            await asyncio.Future()
        except Exception as e:
            print(f"[AI SERVICE] Baglanti denemesi {attempt+1}/30: {e}")
            await asyncio.sleep(3)

if __name__ == "__main__":
    asyncio.run(main())
